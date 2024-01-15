# 111111111111111111111111111111111111

names = ["ana", "Andre", "Ana", "Hide", "Anna"]


def palindrome(list):
    for i in range(len(list)):
        if list[i].lower() == list[i].lower()[::-1]:
            print(f"{list[i]} is a palindrome.")
        else:
            print(f"{list[i]} is not a palindrome.")


palindrome(names)

# 222222222222222222222222222222222


def find_alphanumeric_words(input_str):
    words = input_str.split()

    alphanumeric_words = [
        word
        for word in words
        if any(char.isalpha() for char in word) and any(char.isdigit() for char in word)
    ]

    return alphanumeric_words


input_str = "Emma25 is Data scientist50 and AI Expert"
result = find_alphanumeric_words(input_str)

print("Words with both alphabets and numbers:", result)

# 444444444444444444444444444444444444444

import docx

doc = docx.Document()

doc.add_paragraph("André Martins")
doc.add_paragraph("Hobbies: Video Games")
doc.add_paragraph("Computer Programmer")
doc.add_picture("C:/Users/andre/Anu/Imagens/Fotos Importantes/DSC_0177.JPG")
doc.save("My_document.docx")
